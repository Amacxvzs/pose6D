from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOC_PATH = Path.home() / "Desktop" / "论文_实验图版.docx"
PROJECT_ROOT = Path(r"D:\1.sjcl\pose6d")
FIG_DIR = PROJECT_ROOT / "outputs"


def set_font(run, size: float | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)


def insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        run = new_para.add_run(text)
        set_font(run)
    return new_para


def remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def find_para(doc: Document, pred):
    for paragraph in doc.paragraphs:
        if pred(paragraph):
            return paragraph
    return None


def paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def add_image_after_caption(caption_para, image: Path, width_in: float) -> None:
    pic_p = insert_paragraph_after(caption_para)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(image), width=Inches(width_in))


def replace_placeholder(doc: Document, marker: str, caption: str, image: Path, width_in: float) -> None:
    para = find_para(doc, lambda p: marker in paragraph_text(p))
    if para is None:
        return
    para.text = caption
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        set_font(run, 10.5)
    # Insert image only if the next paragraph is not already an image.
    add_image_after_caption(para, image, width_in)


def remove_caption_and_following_image(doc: Document, caption_prefix: str) -> None:
    paras = list(doc.paragraphs)
    for i, para in enumerate(paras):
        if paragraph_text(para).startswith(caption_prefix):
            if i + 1 < len(paras) and has_drawing(paras[i + 1]):
                remove_paragraph(paras[i + 1])
            remove_paragraph(para)
            return


def rename_caption(doc: Document, old_prefix: str, new_text: str) -> None:
    para = find_para(doc, lambda p: paragraph_text(p).startswith(old_prefix))
    if para is None:
        return
    para.text = new_text
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        set_font(run, 10.5)


def main() -> None:
    doc = Document(str(DOC_PATH))
    pipeline = FIG_DIR / "thesis_figures" / "fig_rgbd_two_stage_pipeline.png"
    difficulty = FIG_DIR / "thesis_figures" / "fig_difficulty_summary.png"
    failure = FIG_DIR / "failure_review" / "fig_failure_cases.png"

    # Remove duplicate figures added later in the document.
    remove_caption_and_following_image(doc, "图 4 二阶段 RGB-D")
    remove_caption_and_following_image(doc, "图 6 测试集不同遮挡难度")

    # Fill the original placeholders.
    replace_placeholder(doc, "【图 1 整体网络架构图】", "图 1 二阶段 RGB-D 堆叠零件位姿估计流程", pipeline, 5.9)
    replace_placeholder(doc, "【图 2 不同遮挡程度的性能对比】", "图 2 测试集不同遮挡难度下的处理结果", difficulty, 5.6)
    replace_placeholder(doc, "【图 3 失败案例占比饼图】", "图 3 低置信失败案例示例", failure, 5.6)

    # Renumber the additional experimental figures.
    rename_caption(doc, "图 5 二阶段遮挡恢复", "图 4 二阶段遮挡恢复前后有效位姿输出统计")
    rename_caption(doc, "图 7 测试集不同后处理策略", "图 5 测试集不同后处理策略消融对比")

    doc.save(str(DOC_PATH))
    print(DOC_PATH)


if __name__ == "__main__":
    main()

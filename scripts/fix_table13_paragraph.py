from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOC_PATH = Path.home() / "Desktop" / "论文.docx"


def set_font(paragraph):
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def main():
    doc = Document(str(DOC_PATH))
    replacement = (
        "二阶段遮挡恢复结果如下：训练集总目标 479 个，第一阶段可用 356 个，二阶段候选 101 个且全部恢复，最终可用 457 个，可用率 95.4%，RMSE 均值 6.66 mm，IoU 均值 0.732；"
        "验证集总目标 70 个，第一阶段可用 49 个，二阶段候选 17 个且全部恢复，最终可用 66 个，可用率 94.3%，RMSE 均值 7.18 mm，IoU 均值 0.714；"
        "测试集总目标 57 个，第一阶段可用 44 个，二阶段候选 9 个且全部恢复，最终可用 53 个，可用率 93.0%，RMSE 均值 6.30 mm，IoU 均值 0.737。"
        "由此可见，二阶段遮挡恢复显著提高了堆叠场景下的有效位姿输出率。"
    )
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("由表 13 可见"):
            paragraph.text = replacement
            set_font(paragraph)
            break
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("表 13"):
            paragraph.text = "二阶段遮挡恢复结果"
            set_font(paragraph)
            break
    doc.save(str(DOC_PATH))
    print(DOC_PATH)


if __name__ == "__main__":
    main()

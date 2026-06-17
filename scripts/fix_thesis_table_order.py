from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOC_PATH = Path.home() / "Desktop" / "论文.docx"


def paragraph_text(p) -> str:
    return "".join((t.text or "") for t in p._p.iter(qn("w:t"))).strip()


def table_text(tbl) -> str:
    return "|".join((t.text or "") for t in tbl.iter(qn("w:t")) if t.text)


def find_paragraph(doc: Document, startswith: str):
    for p in doc.paragraphs:
        if paragraph_text(p).startswith(startswith):
            return p
    raise ValueError(f"paragraph not found: {startswith}")


def find_table(doc: Document, contains: str):
    for table in doc.tables:
        if contains in table_text(table._tbl):
            return table
    raise ValueError(f"table not found: {contains}")


def move_table_after_paragraph(table, paragraph) -> None:
    tbl = table._tbl
    parent = tbl.getparent()
    parent.remove(tbl)
    paragraph._p.addnext(tbl)


def main() -> None:
    doc = Document(str(DOC_PATH))
    cap13 = find_paragraph(doc, "表 13")
    cap14 = find_paragraph(doc, "表 14 测试集")
    table13 = find_table(doc, "总目标数|第一阶段可用")
    table14 = find_table(doc, "无遮挡或分离|22|22")
    move_table_after_paragraph(table13, cap13)
    move_table_after_paragraph(table14, cap14)
    doc.save(str(DOC_PATH))
    print(DOC_PATH)


if __name__ == "__main__":
    main()

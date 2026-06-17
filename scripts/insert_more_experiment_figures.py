from __future__ import annotations

from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOC_PATH = Path.home() / "Desktop" / "论文.docx"
PROJECT_ROOT = Path(r"D:\1.sjcl\pose6d")


def set_font(run, size: float | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)


def insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        run = new_para.add_run(text)
        set_font(run)
    return new_para


def find_contains(doc: Document, needle: str):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(needle)


def add_picture_after(paragraph, image: Path, caption: str, width_in: float = 5.6):
    cap = insert_paragraph_after(paragraph, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_font(run, 10.5)
    pic_p = insert_paragraph_after(cap)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(image), width=Inches(width_in))
    return pic_p


def replace_placeholder_with_picture(doc: Document, marker: str, image: Path, caption: str, width_in: float = 5.6):
    paragraph = find_contains(doc, marker)
    paragraph.text = caption
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        set_font(run, 10.5)
    pic_p = insert_paragraph_after(paragraph)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(image), width=Inches(width_in))


def main() -> None:
    ablation_fig = PROJECT_ROOT / "outputs" / "ablation_experiment" / "fig_ablation_test_strategies.png"
    failure_fig = PROJECT_ROOT / "outputs" / "failure_review" / "fig_failure_cases.png"
    difficulty_fig = PROJECT_ROOT / "outputs" / "thesis_figures" / "fig_difficulty_summary.png"
    for path in [ablation_fig, failure_fig, difficulty_fig]:
        if not path.exists():
            raise FileNotFoundError(path)

    backup = DOC_PATH.with_name(f"论文_实验图备份_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(DOC_PATH, backup)
    doc = Document(str(DOC_PATH))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    if "图 2 测试集不同遮挡难度下的处理结果" not in full_text:
        replace_placeholder_with_picture(
            doc,
            "【图 2 不同遮挡程度的性能对比】",
            difficulty_fig,
            "图 2 测试集不同遮挡难度下的处理结果",
            5.6,
        )

    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "图 3 低置信失败案例示例" not in full_text:
        replace_placeholder_with_picture(
            doc,
            "【图 3 失败案例占比饼图】",
            failure_fig,
            "图 3 低置信失败案例示例",
            5.6,
        )

    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "图 7 测试集不同后处理策略消融对比" not in full_text:
        anchor = find_contains(doc, "二阶段遮挡恢复结果如下")
        add_picture_after(
            anchor,
            ablation_fig,
            "图 7 测试集不同后处理策略消融对比",
            5.6,
        )

    try:
        doc.save(str(DOC_PATH))
        saved = DOC_PATH
    except PermissionError:
        saved = DOC_PATH.with_name("论文_实验图版.docx")
        doc.save(str(saved))
    print(f"backup={backup}")
    print(f"saved={saved}")


if __name__ == "__main__":
    main()

from __future__ import annotations

from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(r"D:\1.sjcl\pose6d")
DOC_PATH = Path.home() / "Desktop" / "论文.docx"
FIG_DIR = PROJECT_ROOT / "outputs" / "thesis_figures"


def set_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_font(run)
    return new_para


def find_paragraph_contains(doc: Document, needle: str):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"paragraph containing {needle!r} not found")


def insert_figure_after(anchor, image_path: Path, caption: str, width_in: float = 5.8):
    cap = insert_paragraph_after(anchor, caption, "Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_font(run, 10.5)

    fig_p = insert_paragraph_after(cap)
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    return fig_p


def remove_existing_caption_blocks(doc: Document) -> None:
    captions = [
        "图 4 二阶段 RGB-D 堆叠零件位姿估计流程",
        "图 5 二阶段遮挡恢复前后有效位姿输出统计",
        "图 6 测试集不同遮挡难度下的处理结果",
    ]
    # Only remove captions on rerun. Existing images are harmless if this script
    # is not rerun often, but removing image paragraphs robustly in mixed Word XML
    # is riskier than avoiding duplicate captions.
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in captions:
            parent = paragraph._element.getparent()
            if parent is not None:
                parent.remove(paragraph._element)


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)
    for name in [
        "fig_rgbd_two_stage_pipeline.png",
        "fig_two_stage_summary.png",
        "fig_difficulty_summary.png",
    ]:
        if not (FIG_DIR / name).exists():
            raise FileNotFoundError(FIG_DIR / name)

    backup = DOC_PATH.with_name(f"论文_插图前备份_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(DOC_PATH, backup)

    doc = Document(str(DOC_PATH))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    if "图 4 二阶段 RGB-D 堆叠零件位姿估计流程" not in full_text:
        anchor = find_paragraph_contains(doc, "该流程避免了后景目标点云被前景点云抢占的问题")
        insert_figure_after(
            anchor,
            FIG_DIR / "fig_rgbd_two_stage_pipeline.png",
            "图 4 二阶段 RGB-D 堆叠零件位姿估计流程",
            width_in=5.9,
        )

    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "图 5 二阶段遮挡恢复前后有效位姿输出统计" not in full_text:
        anchor = find_paragraph_contains(doc, "YOLOv9t 检测模型在测试集上达到 mAP50=0.973")
        insert_figure_after(
            anchor,
            FIG_DIR / "fig_two_stage_summary.png",
            "图 5 二阶段遮挡恢复前后有效位姿输出统计",
            width_in=5.6,
        )

    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "图 6 测试集不同遮挡难度下的处理结果" not in full_text:
        anchor = find_paragraph_contains(doc, "表 14 表明，无遮挡或分离样本的直接可用率")
        insert_figure_after(
            anchor,
            FIG_DIR / "fig_difficulty_summary.png",
            "图 6 测试集不同遮挡难度下的处理结果",
            width_in=5.6,
        )

    doc.save(str(DOC_PATH))
    print(f"backup={backup}")
    print(f"saved={DOC_PATH}")


if __name__ == "__main__":
    main()

from __future__ import annotations

import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path.home() / "Desktop" / "论文_最终实验版.docx"
OUT_DIR = Path.home() / "Desktop" / "Measurement投稿整理包"
MANUSCRIPT = OUT_DIR / "01_Anonymous_Manuscript_Format_Draft.docx"
TITLE_PAGE = OUT_DIR / "02_Title_Page_Template.docx"
HIGHLIGHTS = OUT_DIR / "03_Highlights.docx"
COVER_LETTER = OUT_DIR / "04_Cover_Letter_Template.docx"
FIGURE_DIR = OUT_DIR / "Figures"


TITLE = (
    "RGB-D 6D pose measurement of stacked axisymmetric industrial parts "
    "using YOLOv9, two-stage depth masking, and table-plane constraints"
)

ABSTRACT = (
    "Reliable pose measurement of small stacked industrial parts is difficult because "
    "occlusion, depth mixing, and axial symmetry can produce correct image projections "
    "but physically incorrect orientations. This study develops an RGB-D 6D pose "
    "measurement pipeline combining YOLOv9 detection, object-level depth extraction, "
    "CAD-to-point-cloud registration, two-stage foreground masking, and table-plane "
    "normal constraints. A manually labelled dataset containing 200 training, 30 "
    "validation, and 30 test images was expanded to 698 training images through "
    "model-assisted pseudo-labelling and manual review of low-confidence cases. After "
    "correcting one missing test annotation, the detector achieved 99.36% precision, "
    "100.00% recall, 99.50% mAP50, and 89.34% mAP50-95 on 60 test instances. In the "
    "pose stage, 41 targets were accepted directly and 15 occluded targets were recovered "
    "by removing foreground depth regions, giving 56 position-valid poses with a mean "
    "ICP RMSE of 6.60 mm and a mean reprojection IoU of 0.740. RANSAC table-plane fitting "
    "had an average residual RMSE of approximately 1.93 mm on the test set. Constraining "
    "part normals to remain within 20 degrees of the fitted table normal corrected six "
    "degenerate orientations and produced 54 strictly valid poses, corresponding to a "
    "90.0% strict validity rate. The results demonstrate that measurement-scene geometry "
    "can provide an effective physical constraint for pose estimation of axisymmetric parts."
)

KEYWORDS = (
    "6D pose measurement; RGB-D sensing; YOLOv9; point-cloud registration; "
    "occlusion handling; table-plane constraint; axisymmetric parts"
)

HIGHLIGHT_LINES = [
    "A YOLOv9-based RGB-D pipeline measures poses of stacked industrial parts.",
    "Reviewed pseudo-labels expand training data without changing the test set.",
    "Two-stage depth masking recovers rear-object poses under occlusion.",
    "Table-plane constraints correct degenerate normals of axisymmetric parts.",
]


HEADING_MAP = {
    "一、引言": "1. Introduction",
    "二、相关工作": "2. Related work",
    "三、本文提出的方法": "3. Proposed measurement method",
    "四、实验与结果分析": "4. Experiments and results",
    "五、结论与展望": "5. Conclusions",
    "参考文献（IEEE 格式，统一 DOI）": "References",
    "2.1 6D 物体位姿估计研究": "2.1 6D object pose estimation",
    "2.2 基于 Transformer 的 6D 位姿估计": "2.2 Transformer-based 6D pose estimation",
    "2.3 YOLO 系列算法在工业检测中的应用": "2.3 YOLO methods in industrial inspection",
    "2.4 堆叠物体位姿估计研究": "2.4 Pose estimation of stacked objects",
    "3.1 整体网络架构": "3.1 Overall architecture",
    "3.2 适配后的遮挡感知注意力模块": "3.2 Occlusion-aware attention adaptation",
    "3.3 多任务输出分支与 EPnP 联合求解": "3.3 Multi-task outputs and EPnP solution",
    "3.3.1 多任务输出分支": "3.3.1 Multi-task output branches",
    "3.3.2 角点投影损失": "3.3.2 Corner-projection loss",
    "3.3.3 四元数辅助损失": "3.3.3 Auxiliary quaternion loss",
    "3.3.4 EPnP 联合求解 6D 位姿": "3.3.4 EPnP-based 6D pose solution",
    "3.4 堆叠几何一致性后处理": "3.4 Geometric consistency post-processing",
    "3.5 基于 RGB-D 点云的二阶段遮挡恢复策略": "3.5 Two-stage RGB-D occlusion recovery",
    "4.1 数据集构建": "4.1 Dataset construction",
    "4.1.1 自建工业数据集": "4.1.1 In-house industrial dataset",
    "4.1.2 位姿真值与角点标签获取": "4.1.2 Pose ground truth and corner labels",
    "4.1.3 跨数据集验证": "4.1.3 Cross-dataset validation",
    "4.1.4 跨类别泛化验证": "4.1.4 Cross-category generalization",
    "4.1.5 标签敏感性分析": "4.1.5 Label-sensitivity analysis",
    "4.2 实验环境": "4.2 Experimental setup",
    "4.3 评价指标": "4.3 Evaluation metrics",
    "4.4 训练细节与基线方法说明": "4.4 Training details and baselines",
    "4.4.1 训练超参数设置": "4.4.1 Training hyperparameters",
    "4.4.2 基线方法说明": "4.4.2 Baseline methods",
    "4.5 对比实验": "4.5 Comparative experiments",
    "4.6 消融实验": "4.6 Ablation experiments",
    "4.6.1 卷积核尺寸消融实验": "4.6.1 Convolution-kernel ablation",
    "4.6.2 关键点选择对比实验": "4.6.2 Keypoint-selection comparison",
    "4.6.3 四元数辅助损失消融实验": "4.6.3 Quaternion-loss ablation",
    "4.7 推理重复性分析": "4.7 Repeatability analysis",
    "4.8 后处理效果统计与阈值消融": "4.8 Post-processing and threshold analysis",
    "4.9 端到端延迟分析": "4.9 End-to-end latency",
    "4.10 不同遮挡程度的鲁棒性分析": "4.10 Robustness under different occlusion levels",
    "4.11 失败案例分析": "4.11 Failure analysis",
    "4.12 当前系统实现与阶段性实验结果": "4.12 Verified RGB-D system results",
    "4.13 工作台平面约束与方向优化": "4.13 Table-plane constraint and orientation correction",
}


def set_font(run, name: str = "Times New Roman", size: float = 12, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def replace_text(paragraph, text: str, size=12, bold=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def remove_toc_paragraphs(doc: Document):
    for paragraph in list(doc.paragraphs):
        if paragraph.style.name.lower().startswith("toc"):
            paragraph._element.getparent().remove(paragraph._element)


def format_manuscript():
    doc = Document(str(SOURCE))
    remove_toc_paragraphs(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.columns = 1
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = None

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 12)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    replace_text(doc.paragraphs[0], TITLE, size=16, bold=True)
    doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    replace_text(doc.paragraphs[1], "Abstract\n" + ABSTRACT, size=11)
    doc.paragraphs[1].paragraph_format.line_spacing = 1.15
    replace_text(doc.paragraphs[2], "Keywords: " + KEYWORDS, size=11)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in HEADING_MAP:
            replace_text(paragraph, HEADING_MAP[text], size=14 if paragraph.style.name == "Heading 2" else 12, bold=True)
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.text.strip().startswith("图 "):
            number, _, rest = paragraph.text.strip().partition(" ")
            replace_text(paragraph, "Figure " + rest, size=10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
        elif paragraph.text.strip().startswith("表 "):
            replace_text(paragraph, paragraph.text.strip().replace("表 ", "Table ", 1), size=10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("["):
            paragraph.paragraph_format.left_indent = Cm(0.75)
            paragraph.paragraph_format.first_line_indent = Cm(-0.75)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(3)
            for run in paragraph.runs:
                set_font(run, size=10)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_font(run, size=9, bold=(row_index == 0))

    doc.core_properties.author = "Anonymous"
    doc.core_properties.last_modified_by = "Anonymous"
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Measurement submission format draft"
    doc.save(str(MANUSCRIPT))


def make_title_page():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    set_font(run, size=16, bold=True)
    fields = [
        ("Authors", "[Author 1], [Author 2], ..."),
        ("Affiliations", "[Department, Institution, City, Postal code, Country]"),
        ("Corresponding author", "[Full name]"),
        ("Email", "[Institutional email]"),
        ("Telephone", "[Optional]"),
        ("ORCID", "[ORCID iD for each author, if available]"),
        ("Declaration of interests", "The authors declare no competing interests. [Confirm/edit]"),
        ("Funding", "[Funding agency and grant number, or: This research received no external funding.]"),
        ("Data availability", "[Repository/DOI or a justified availability statement]"),
        ("Author contributions", "[Use CRediT roles]"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        r = p.add_run(label + ": ")
        set_font(r, size=12, bold=True)
        r = p.add_run(value)
        set_font(r, size=12)
        p.paragraph_format.space_after = Pt(8)
    doc.save(str(TITLE_PAGE))


def make_highlights():
    doc = Document()
    title = doc.add_paragraph()
    r = title.add_run("Highlights")
    set_font(r, size=14, bold=True)
    for line in HIGHLIGHT_LINES:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_font(r, size=12)
    doc.save(str(HIGHLIGHTS))


def make_cover_letter():
    doc = Document()
    paragraphs = [
        "Dear Editor-in-Chief,",
        (
            f"We submit the manuscript entitled “{TITLE}” for consideration as a Research Paper "
            "in Measurement."
        ),
        (
            "The manuscript addresses measurement of 6D poses for stacked axisymmetric industrial "
            "parts using synchronized RGB-D sensing. Its measurement-science contribution is a "
            "physically constrained evaluation and correction framework that combines object-level "
            "depth measurement, CAD registration, occlusion-aware recovery, and RANSAC table-plane "
            "normal estimation."
        ),
        (
            "The work reports detector accuracy, point-cloud registration RMSE, reprojection IoU, "
            "table-plane fitting residuals, and strict pose-validity rates. The manuscript is original, "
            "has not been published previously, and is not under consideration elsewhere. "
            "[Confirm before submission.]"
        ),
        "All authors have approved the manuscript and agree with its submission. [Confirm.]",
        "Sincerely,",
        "[Corresponding author name]",
        "[Institution]",
        "[Email]",
    ]
    for index, text in enumerate(paragraphs):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, size=12)
        p.paragraph_format.space_after = Pt(10)
    doc.save(str(COVER_LETTER))


def extract_figures():
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(MANUSCRIPT))
    index = 0
    for paragraph in doc.paragraphs:
        for drawing in paragraph._p.xpath(".//w:drawing"):
            blips = drawing.xpath(".//a:blip")
            if not blips:
                continue
            rel_id = blips[0].get(qn("r:embed"))
            if not rel_id:
                continue
            part = doc.part.related_parts[rel_id]
            index += 1
            suffix = Path(part.partname).suffix or ".png"
            (FIGURE_DIR / f"Figure_{index:02d}{suffix}").write_bytes(part.blob)


def make_checklist():
    text = """# Measurement Submission Checklist

## Completed in this package

- Anonymous single-column Word manuscript format draft.
- Separate title-page template.
- Separate Highlights file with four bullet points.
- Separate cover-letter template.
- Figures extracted into individual files.
- English title, abstract (under 250 words), and 1-7 keywords.
- Numbered section hierarchy and numbered reference list.

## Mandatory before submission

- Translate the entire Chinese body, all tables, and all figure captions into polished English.
- Remove or replace all unverified claims involving 8 classes, 7200 images,
  LINEMOD-Occlusion, ADD/ADD-S=92.3%, 27 FPS, and similar planned results.
- Condense the manuscript to Measurement's research-paper length expectation
  (normally no more than 30 pages).
- Add real author names, affiliations, ORCID identifiers, corresponding-author
  details, CRediT roles, funding, data availability, and competing-interest text.
- Provide a reproducible calibration and uncertainty analysis:
  camera calibration uncertainty, depth repeatability, table-plane fitting
  uncertainty, and pose repeatability across repeated captures.
- Upload editable figures or high-resolution TIFF/JPEG/EPS/PDF files separately.
- Confirm every citation and DOI against the original publication.
- Add a data/code repository or a justified data-availability statement.
- Verify the final English manuscript with plagiarism, grammar, and reference checks.

## Measurement fit

The paper should foreground the measurement contribution rather than present
the work only as an AI detection application. Emphasize:

- the measured quantities and traceable coordinate systems;
- calibration, repeatability, uncertainty, and error propagation;
- the physical table-plane constraint and why it improves measurement validity;
- reproducible evaluation criteria for axisymmetric objects.
"""
    (OUT_DIR / "05_Submission_Checklist.md").write_text(text, encoding="utf-8")


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True)
    format_manuscript()
    make_title_page()
    make_highlights()
    make_cover_letter()
    extract_figures()
    make_checklist()
    print(OUT_DIR)


if __name__ == "__main__":
    main()

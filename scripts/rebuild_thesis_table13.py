from __future__ import annotations

from pathlib import Path
import csv

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(r"D:\1.sjcl\pose6d")
DOC_PATH = Path.home() / "Desktop" / "论文.docx"


def read_csv(path: Path) -> list[dict]:
    with path.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def pct(value: str) -> str:
    return f"{float(value) * 100:.1f}%"


def f2(value: str) -> str:
    return f"{float(value):.2f}"


def f3(value: str) -> str:
    return f"{float(value):.3f}"


def paragraph_text(p) -> str:
    return "".join((t.text or "") for t in p._p.iter(qn("w:t"))).strip()


def table_header(table) -> str:
    if not table.rows:
        return ""
    return " / ".join(cell.text for cell in table.rows[0].cells)


def remove_bad_table13(doc: Document) -> None:
    for table in list(doc.tables):
        header = table_header(table)
        if "总目标数" in header and "第一阶段可用" in header:
            tbl = table._tbl
            parent = tbl.getparent()
            if parent is not None:
                parent.remove(tbl)


def find_cap13(doc: Document):
    for p in doc.paragraphs:
        if paragraph_text(p).startswith("表 13"):
            return p
    raise ValueError("表 13 caption not found")


def set_run_font(run, size: float = 8.5, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def insert_table_after(doc: Document, paragraph, data: list[list[str]]) -> None:
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    for style_name in ("Table Grid", "网格型"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(data):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(text)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, 8.5, i == 0)
    tbl = table._tbl
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)
    paragraph._p.addnext(tbl)


def main() -> None:
    doc = Document(str(DOC_PATH))
    remove_bad_table13(doc)
    cap13 = find_cap13(doc)
    two_stage = {r["split"]: r for r in read_csv(PROJECT_ROOT / "outputs" / "two_stage_summary.csv")}
    data = [["数据集", "总目标数", "第一阶段可用", "二阶段候选", "二阶段恢复", "最终可用", "可用率", "RMSE均值/mm", "IoU均值"]]
    for split in ["train", "val", "test"]:
        r = two_stage[split]
        data.append(
            [
                {"train": "训练集", "val": "验证集", "test": "测试集"}[split],
                r["total_original"],
                r["stage1_count"],
                r["stage2_candidate_count"],
                r["stage2_recovered"],
                r["final_usable"],
                pct(r["final_usable_rate"]),
                f2(r["rmse_mean_mm"]),
                f3(r["final_iou_mean"]),
            ]
        )
    insert_table_after(doc, cap13, data)
    doc.save(str(DOC_PATH))
    print(DOC_PATH)


if __name__ == "__main__":
    main()

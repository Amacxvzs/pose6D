from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOC_PATH = Path.home() / "Desktop" / "论文.docx"


def set_run_font(run) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def paragraph_text(paragraph) -> str:
    return "".join((t.text or "") for t in paragraph._p.iter(qn("w:t"))).strip()


def table_header(table) -> str:
    if not table.rows:
        return ""
    return " / ".join(cell.text for cell in table.rows[0].cells)


def remove_table13(doc: Document) -> None:
    for table in list(doc.tables):
        header = table_header(table)
        if "总目标数" in header and "第一阶段可用" in header:
            tbl = table._tbl
            parent = tbl.getparent()
            if parent is not None:
                parent.remove(tbl)


def find_paragraph(doc: Document, startswith: str):
    for p in doc.paragraphs:
        if paragraph_text(p).startswith(startswith):
            return p
    raise ValueError(startswith)


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    run = new_para.add_run(text)
    set_run_font(run)
    return new_para


def main() -> None:
    doc = Document(str(DOC_PATH))
    remove_table13(doc)
    cap13 = find_paragraph(doc, "表 13")
    cap13.text = "二阶段遮挡恢复结果如下："
    for run in cap13.runs:
        set_run_font(run)

    # Avoid duplicate insertion.
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "训练集：总目标 479 个" not in full_text:
        cur = cap13
        for text in [
            "训练集：总目标 479 个，第一阶段可用 356 个，二阶段候选 101 个且全部恢复，最终可用 457 个，可用率 95.4%，RMSE 均值 6.66 mm，IoU 均值 0.732。",
            "验证集：总目标 70 个，第一阶段可用 49 个，二阶段候选 17 个且全部恢复，最终可用 66 个，可用率 94.3%，RMSE 均值 7.18 mm，IoU 均值 0.714。",
            "测试集：总目标 57 个，第一阶段可用 44 个，二阶段候选 9 个且全部恢复，最终可用 53 个，可用率 93.0%，RMSE 均值 6.30 mm，IoU 均值 0.737。",
        ]:
            cur = insert_paragraph_after(cur, text)
    doc.save(str(DOC_PATH))
    print(DOC_PATH)


if __name__ == "__main__":
    main()

from __future__ import annotations

from datetime import datetime
from pathlib import Path
import csv
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(r"D:\1.sjcl\pose6d")
DOC_PATH = Path.home() / "Desktop" / "论文.docx"


def read_csv(path: Path) -> list[dict]:
    with path.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def pct(value: str) -> str:
    return f"{float(value) * 100:.1f}%"


def f2(value: str) -> str:
    return f"{float(value):.2f}"


def f3(value: str) -> str:
    return f"{float(value):.3f}"


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def insert_table_after(paragraph, data: list[list[str]]):
    table = paragraph._parent.add_table(rows=len(data), cols=len(data[0]), width=Inches(6.5))
    paragraph._p.addnext(table._tbl)
    for style_name in ("Table Grid", "网格型"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(data):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(text)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=8.5, bold=(i == 0))
    return table


def find_index(doc: Document, exact: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == exact:
            return i
    raise ValueError(f"paragraph not found: {exact}")


def add_method_section(doc: Document) -> None:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "3.5 基于 RGB-D 点云的二阶段遮挡恢复策略" in full_text:
        return
    idx4 = find_index(doc, "四、实验与结果分析")
    cur = insert_paragraph_after(
        doc.paragraphs[idx4 - 1],
        "3.5 基于 RGB-D 点云的二阶段遮挡恢复策略",
        "Heading 3",
    )
    paragraphs = [
        "在当前实验实现中，考虑到目标零件为近似轴对称的盘类零件，仅依赖 8 个 AABB 角点的平面内旋转约束会受到对称性影响。为提高真实堆叠场景下的位姿稳定性，本文在 YOLOv9 检测结果之后引入 RGB-D 点云约束：首先依据相机内参和裁剪元数据，将 720×720 裁剪图上的检测框还原至 1280×720 彩色全图坐标，再从同步采集的 16 位深度图中提取目标局部点云。",
        "点云提取阶段以检测框为初始 ROI，利用深度有效范围与近邻深度窗口筛选前景点，并通过连通域分析去除背景点和孤立噪声。对于盘类零件，初始位姿不再强制估计不可观测的绕对称轴 yaw 角，而是由点云质心确定平移量，由 PCA 平面法向确定零件法向；随后使用 CAD 模型采样点云与观测点云进行 ICP 配准，得到几何一致的 6D 位姿。",
        "针对相互遮挡的堆叠样本，本文进一步设计前景优先的二阶段恢复流程。第一阶段根据深度近远关系、检测框重叠度、ICP 适配度、重投影误差和点云数量对每个目标进行质量评估，将样本划分为可靠、可用和低置信三类；第二阶段对被前景零件遮挡的后景候选目标，根据 blocked_by_det_ids 将前景遮挡区域从深度图中剔除，然后重新提取后景点云并再次执行初始位姿估计与 ICP。该流程避免了后景目标点云被前景点云抢占的问题，适合本文的盘类小零件堆叠实验。",
    ]
    for text in paragraphs:
        cur = insert_paragraph_after(cur, text, "Normal")


def add_experiment_section(doc: Document) -> None:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "4.12 当前系统实现与阶段性实验结果" in full_text:
        return

    two_stage = {r["split"]: r for r in read_csv(PROJECT_ROOT / "outputs" / "two_stage_summary.csv")}
    difficulty = read_csv(PROJECT_ROOT / "outputs" / "pose_difficulty_summary.csv")

    idx_conclusion = find_index(doc, "五、结论与展望")
    cur = insert_paragraph_after(
        doc.paragraphs[idx_conclusion - 1],
        "4.12 当前系统实现与阶段性实验结果",
        "Heading 3",
    )
    cur = insert_paragraph_after(
        cur,
        "为验证上述流程的工程可行性，本文以盘类打印零件为目标对象，使用 Astra Pro Plus 深度相机完成 RGB-D 图像采集，并基于 CAD 模型生成目标几何先验。当前实验共采集 260 组样本，按训练集 200 组、验证集 30 组、测试集 30 组进行划分；YOLOv9t 检测模型在测试集上达到 mAP50=0.973、mAP50-95=0.868。在位姿阶段，系统依次执行检测、RGB-D 局部点云提取、PCA 初始姿态估计、CAD-ICP 精修、质量评估和二阶段遮挡恢复。",
        "Normal",
    )

    caption = insert_paragraph_after(cur, "表 13 二阶段遮挡恢复后的位姿估计结果", "Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table13 = [["数据集", "总目标数", "第一阶段可用", "二阶段候选", "二阶段恢复", "最终可用", "可用率", "RMSE均值/mm", "IoU均值"]]
    for split in ["train", "val", "test"]:
        r = two_stage[split]
        table13.append(
            [
                {"train": "训练集", "val": "验证集", "test": "测试集"}[split],
                r["total_original"],
                r["stage1_count"],
                r["stage2_candidate_count"],
                r["stage2_recovered"],
                r["final_usable"],
                pct(r["final_usable_rate"]),
                f2(r["rmse_mean_mm"]),
                f3(r["final_iou_mean"]),
            ]
        )
    insert_table_after(caption, table13)

    cur = insert_paragraph_after(
        caption,
        "由表 13 可见，二阶段遮挡恢复显著提高了堆叠场景下的有效位姿输出率。测试集共有 57 个检测目标，其中第一阶段可直接使用 44 个；经过前景区域剔除和后景点云重提取后，9 个遮挡候选目标全部恢复成功，最终可用目标数提升至 53 个，可用率达到 93.0%，测试集平均 ICP RMSE 为 6.30 mm，平均重投影 IoU 为 0.737。",
        "Normal",
    )

    caption2 = insert_paragraph_after(cur, "表 14 测试集不同遮挡难度下的阶段性结果", "Normal")
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    labels = {
        "clear_single_or_separated": "无遮挡或分离",
        "partial_overlap": "轻微重叠",
        "occluded_stack": "堆叠遮挡",
    }
    table14 = [["难度类型", "目标数", "直接可用", "二阶段候选", "低置信", "直接可用率", "RMSE均值/mm", "IoU均值"]]
    for r in [x for x in difficulty if x["split"] == "test"]:
        table14.append(
            [
                labels[r["difficulty"]],
                r["total"],
                r["accepted"],
                r["stage2_candidates"],
                r["low_confidence"],
                pct(r["accept_rate"]),
                "-" if r["difficulty"] == "occluded_stack" else f2(r["rmse_mean_mm"]),
                "-" if r["difficulty"] == "occluded_stack" else f3(r["final_iou_mean"]),
            ]
        )
    insert_table_after(caption2, table14)

    cur = insert_paragraph_after(
        caption2,
        "表 14 表明，无遮挡或分离样本的直接可用率为 100.0%，轻微重叠样本的直接可用率为 91.7%；而堆叠遮挡样本在第一阶段不直接进入主统计，而是作为二阶段遮挡恢复对象处理。该结果说明，单纯依赖检测框局部深度容易在前后零件重叠处混入错误点云，而前景优先估计与深度区域剔除能够有效恢复后景目标的点云质量。",
        "Normal",
    )
    insert_paragraph_after(
        cur,
        "需要说明的是，本文当前实验对象为近似轴对称的盘类零件，因此绕零件法向的平面内旋转角在视觉上并不唯一。后续定量评价更适合采用中心位置误差、法向方向误差、重投影 IoU、ICP RMSE 以及抓取可用率等指标，而不宜将对称轴 yaw 角作为主要评价项。",
        "Normal",
    )


def add_conclusion_update(doc: Document) -> None:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "系统已完成从数据采集、YOLOv9t 检测训练、RGB-D 点云提取" in full_text:
        return
    idx = find_index(doc, "五、结论与展望")
    insert_paragraph_after(
        doc.paragraphs[idx],
        "基于当前实物采集实验，系统已完成从数据采集、YOLOv9t 检测训练、RGB-D 点云提取、CAD-ICP 精修到二阶段遮挡恢复的完整闭环。实验结果显示，测试集最终可用位姿为 53/57，可用率为 93.0%，平均 RMSE 为 6.30 mm，说明本文方法在盘类工业零件的轻微重叠与部分堆叠场景中具有较好的工程可行性。",
        "Normal",
    )


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)
    backup_path = DOC_PATH.with_name(f"论文_备份_{datetime.now():%Y%m%d_%H%M%S}.docx")
    shutil.copy2(DOC_PATH, backup_path)

    doc = Document(str(DOC_PATH))
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    add_method_section(doc)
    add_experiment_section(doc)
    add_conclusion_update(doc)

    doc.save(str(DOC_PATH))
    print(f"backup={backup_path}")
    print(f"saved={DOC_PATH}")


if __name__ == "__main__":
    main()

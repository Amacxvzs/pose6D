from __future__ import annotations

import csv
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
INPUT_DOC = Path.home() / "Desktop" / "论文_实验图版.docx"
OUTPUT_DOC = Path.home() / "Desktop" / "论文_最终实验版.docx"
FIG_DIR = PROJECT_ROOT / "outputs" / "thesis_figures"


def read_csv(path: Path) -> list[dict]:
    with path.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")


def font(size: int, bold: bool = False):
    path = Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else FONT_PATH
    return ImageFont.truetype(str(path), size)


def chart_canvas(title: str, y_min: float, y_max: float):
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((800, 42), title, font=font(42, True), fill="#18384f", anchor="ma")
    left, top, right, bottom = 150, 135, 1510, 760
    draw.line((left, top, left, bottom), fill="#333333", width=3)
    draw.line((left, bottom, right, bottom), fill="#333333", width=3)
    for index in range(6):
        value = y_min + (y_max - y_min) * index / 5
        y = bottom - (bottom - top) * index / 5
        draw.line((left, y, right, y), fill="#dddddd", width=2)
        draw.text((left - 20, y), f"{value:.0f}", font=font(24), fill="#444444", anchor="rm")
    return image, draw, (left, top, right, bottom)


def save_chart(image: Image.Image, name: str) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    path = FIG_DIR / name
    image.save(path, dpi=(240, 240))
    return path


def generate_detection_figure() -> Path:
    labels = ["Precision", "Recall", "mAP50", "mAP50-95"]
    baseline = [94.87, 98.33, 98.90, 88.39]
    enhanced = [99.36, 100.00, 99.50, 89.34]
    image, draw, (left, top, right, bottom) = chart_canvas("补标修正测试集上的检测性能对比", 84, 101)
    group_width = (right - left) / len(labels)
    bar_width = 90
    for index, label in enumerate(labels):
        center = left + group_width * (index + 0.5)
        for offset, value, color in ((-52, baseline[index], "#7798c8"), (52, enhanced[index], "#4c9f70")):
            x1 = center + offset - bar_width / 2
            x2 = center + offset + bar_width / 2
            y = bottom - (value - 84) / 17 * (bottom - top)
            draw.rectangle((x1, y, x2, bottom), fill=color)
            draw.text((center + offset, y - 10), f"{value:.2f}", font=font(23), fill="#222222", anchor="ms")
        draw.text((center, bottom + 28), label, font=font(26), fill="#222222", anchor="ma")
    draw.rectangle((930, 88, 975, 118), fill="#7798c8")
    draw.text((990, 103), "原始 200 张训练集", font=font(23), fill="#222222", anchor="lm")
    draw.rectangle((1240, 88, 1285, 118), fill="#4c9f70")
    draw.text((1300, 103), "扩展 698 张训练集", font=font(23), fill="#222222", anchor="lm")
    return save_chart(image, "fig_detection_model_comparison.png")


def generate_constraint_figure() -> Path:
    two_stage = {row["split"]: row for row in read_csv(PROJECT_ROOT / "outputs" / "two_stage_summary.csv")}
    constrained = {}
    for split in ("train", "val", "test"):
        rows = read_csv(
            PROJECT_ROOT
            / "outputs"
            / "two_stage_poses_table_constrained"
            / split
            / "two_stage_poses.csv"
        )
        constrained[split] = {
            "valid": sum(int(row["plane_pose_valid"]) for row in rows),
            "corrected": sum(int(row["normal_constrained"]) for row in rows),
        }

    labels = ["训练集", "验证集", "测试集"]
    splits = ["train", "val", "test"]
    position = [int(two_stage[s]["final_usable"]) for s in splits]
    strict = [constrained[s]["valid"] for s in splits]
    corrected = [constrained[s]["corrected"] for s in splits]
    maximum = 540
    image, draw, (left, top, right, bottom) = chart_canvas("工作台平面约束前后的有效位姿统计", 0, maximum)
    group_width = (right - left) / 3
    bar_width = 120
    for index, label in enumerate(labels):
        center = left + group_width * (index + 0.5)
        for offset, value, color in ((-70, position[index], "#4c9f70"), (70, strict[index], "#245b7a")):
            x1 = center + offset - bar_width / 2
            x2 = center + offset + bar_width / 2
            y = bottom - value / maximum * (bottom - top)
            draw.rectangle((x1, y, x2, bottom), fill=color)
            draw.text((center + offset, y - 10), str(value), font=font(25), fill="#222222", anchor="ms")
        draw.text((center + 70, bottom - strict[index] / maximum * (bottom - top) - 44), f"修正 {corrected[index]}", font=font(21), fill="#444444", anchor="ms")
        draw.text((center, bottom + 28), label, font=font(28), fill="#222222", anchor="ma")
    draw.rectangle((980, 88, 1025, 118), fill="#4c9f70")
    draw.text((1040, 103), "位置可用", font=font(23), fill="#222222", anchor="lm")
    draw.rectangle((1190, 88, 1235, 118), fill="#245b7a")
    draw.text((1250, 103), "方向约束后严格可用", font=font(23), fill="#222222", anchor="lm")
    return save_chart(image, "fig_table_constraint_summary.png")


def generate_two_stage_figure() -> Path:
    rows = read_csv(PROJECT_ROOT / "outputs" / "two_stage_summary.csv")
    labels = ["训练集", "验证集", "测试集"]
    stage1 = [int(row["stage1_count"]) for row in rows]
    stage2 = [int(row["stage2_recovered"]) for row in rows]
    skip = [int(row["skip_or_recapture"]) for row in rows]
    total = [int(row["total_original"]) for row in rows]
    rate = [float(row["final_usable_rate"]) * 100 for row in rows]
    maximum = 540
    image, draw, (left, top, right, bottom) = chart_canvas("二阶段遮挡恢复后的有效位姿输出", 0, maximum)
    group_width = (right - left) / 3
    bar_width = 210
    for index, label in enumerate(labels):
        center = left + group_width * (index + 0.5)
        values = [
            (stage1[index], "#4c9f70"),
            (stage2[index], "#f2b84b"),
            (skip[index], "#d75a4a"),
        ]
        current_bottom = bottom
        for value, color in values:
            height = value / maximum * (bottom - top)
            draw.rectangle((center - bar_width / 2, current_bottom - height, center + bar_width / 2, current_bottom), fill=color)
            current_bottom -= height
        draw.text((center, current_bottom - 10), f"可用率 {rate[index]:.1f}%", font=font(24), fill="#222222", anchor="ms")
        draw.text((center, bottom + 28), label, font=font(28), fill="#222222", anchor="ma")
    legends = [("第一阶段可用", "#4c9f70"), ("二阶段恢复", "#f2b84b"), ("跳过/重拍", "#d75a4a")]
    for index, (label, color) in enumerate(legends):
        x = 780 + index * 260
        draw.rectangle((x, 88, x + 38, 114), fill=color)
        draw.text((x + 50, 101), label, font=font(21), fill="#222222", anchor="lm")
    return save_chart(image, "fig_two_stage_summary_current.png")


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)


def insert_paragraph_before(anchor, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    paragraph = anchor._parent.add_paragraph()
    paragraph._p = new_p
    if style:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, 10.5)
    return paragraph


def insert_picture_before(anchor, image: Path, caption: str, width: float = 5.8):
    caption_p = insert_paragraph_before(anchor, caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.keep_with_next = True
    picture_p = insert_paragraph_before(anchor)
    picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_p.add_run().add_picture(str(image), width=Inches(width))
    return picture_p


def replace_picture_after_caption(doc: Document, caption_text: str, image: Path, width: float = 5.8) -> None:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() != caption_text:
            continue
        for candidate in doc.paragraphs[index + 1 : index + 4]:
            if candidate._p.xpath(".//w:drawing"):
                for child in list(candidate._p):
                    candidate._p.remove(child)
                candidate.alignment = WD_ALIGN_PARAGRAPH.CENTER
                candidate.add_run().add_picture(str(image), width=Inches(width))
                return
    raise ValueError(f"Caption not found: {caption_text}")


def insert_table_before(anchor, data: list[list[str]], widths: list[float]):
    table = anchor._parent.add_table(rows=len(data), cols=len(data[0]), width=Inches(sum(widths)))
    anchor._p.addprevious(table._tbl)
    if anchor._parent.tables:
        table.style = anchor._parent.tables[0].style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, row in enumerate(data):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.width = Inches(widths[col_index])
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, 8.5, bold=(row_index == 0))
    return table


def find_exact(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(text)


def main() -> None:
    if not INPUT_DOC.exists():
        raise FileNotFoundError(INPUT_DOC)
    detection_fig = generate_detection_figure()
    constraint_fig = generate_constraint_figure()
    two_stage_fig = generate_two_stage_figure()
    review_fig = (
        PROJECT_ROOT
        / "outputs"
        / "two_stage_review_table_constrained"
        / "test"
        / "two_stage_contact_sheet_02.png"
    )

    doc = Document(str(INPUT_DOC))
    set_paragraph_text(
        doc.paragraphs[215],
        "为验证上述流程的工程可行性，本文以盘类打印零件为目标对象，使用 Astra Pro Plus 深度相机完成 RGB-D 图像采集，并基于 CAD 模型生成目标几何先验。原始人工标注数据集包含训练集 200 张、验证集 30 张和测试集 30 张；在此基础上，利用旧模型生成伪标注并对低置信样本人工复核，形成 698 张训练图像、1706 个目标框的扩展训练集。补充修正测试集漏标后，改进 YOLOv9t 在 60 个测试实例上获得 Precision=99.36%、Recall=100.00%、mAP50=99.50%、mAP50-95=89.34%，其中 mAP50-95 较原模型的 88.39% 提升 0.95 个百分点。",
    )
    set_paragraph_text(
        doc.paragraphs[219],
        "采用更新后的检测模型重新运行完整 RGB-D 位姿流程后，训练集、验证集和测试集分别得到 504、79 和 60 个检测目标。测试集中第一阶段直接输出 41 个可用位姿，15 个遮挡目标经前景深度区域剔除后全部完成二阶段恢复，另有 4 个低置信目标被判定为跳过或重拍，最终位置可用位姿为 56/60，可用率为 93.3%，平均 ICP RMSE 为 6.60 mm，平均重投影 IoU 为 0.740。结果表明，扩展训练集改善了检测完整性，而二阶段点云恢复保持了堆叠场景中的高有效输出率。",
    )
    set_paragraph_text(
        doc.paragraphs[223],
        "更新后的难度划分结果显示，测试集包含 20 个无遮挡或分离目标、24 个轻微重叠目标和 16 个堆叠遮挡目标。无遮挡目标第一阶段可用率为 100.0%，轻微重叠目标第一阶段可用率为 87.5%；堆叠遮挡目标则进入前景优先的二阶段处理。该结果进一步说明，在前后零件重叠区域，仅依赖局部深度窗口容易混入前景点云，而利用遮挡顺序和前景区域剔除能够恢复后景目标。",
    )
    set_paragraph_text(
        doc.paragraphs[224],
        "盘类零件近似轴对称，绕自身法向的平面内旋转角不可由外圆轮廓唯一确定，因此本文不将对称轴 yaw 角作为主要评价指标。实际验收中进一步发现，仅以相机光轴作为方向参考会产生错误判断，因为相机光轴与工作台法向约存在 24°夹角。为此，本文采用 RANSAC 从每幅深度图拟合工作台平面，并以工作台法向约束零件法向，结合重投影 IoU 和中心误差筛除方向不可靠结果。",
    )
    set_paragraph_text(
        doc.paragraphs[226],
        "基于当前实物采集实验，系统已完成从数据采集、YOLOv9t 检测训练、RGB-D 点云提取、CAD-ICP 精修、二阶段遮挡恢复到工作台平面法向约束的完整闭环。修正测试集漏标后，新检测模型在 60 个实例上达到 Recall=100.00% 和 mAP50=99.50%；位姿阶段获得 56/60 个位置可用结果，经工作台平面约束后有 54/60 个同时满足方向、投影 IoU 和中心误差要求，严格有效率为 90.0%。",
    )

    replace_picture_after_caption(doc, "图 4 二阶段遮挡恢复前后有效位姿输出统计", two_stage_fig)

    # Update the existing difficulty table.
    difficulty_table = doc.tables[13]
    updated_rows = [
        ["无遮挡或分离", "20", "20", "0", "0", "100.0%", "4.79", "0.794"],
        ["轻微重叠", "24", "21", "0", "3", "87.5%", "8.77", "0.686"],
        ["堆叠遮挡", "16", "0", "16", "1", "0.0%", "-", "-"],
    ]
    for row_index, values in enumerate(updated_rows, start=1):
        for col_index, value in enumerate(values):
            difficulty_table.cell(row_index, col_index).text = value
            for paragraph in difficulty_table.cell(row_index, col_index).paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, 8.5)

    conclusion = find_exact(doc, "五、结论与展望")
    heading = insert_paragraph_before(conclusion, "4.13 工作台平面约束与方向优化", "Heading 3")
    insert_paragraph_before(
        conclusion,
        "针对圆盘投影正确但法向明显倾斜的问题，本文在二阶段位姿结果之后增加工作台平面约束。首先将深度图映射至 1280×720 全分辨率彩色坐标系，并排除所有检测框及其邻域；随后采用 RANSAC 拟合工作台平面。训练集、验证集和测试集的平面拟合平均 RMSE 分别约为 1.94 mm、1.90 mm 和 1.93 mm，说明工作台法向估计较为稳定。对于零件法向与工作台法向夹角超过 20°的结果，在保持零件三维中心不变的前提下限制其倾角，并重新计算投影 IoU 与中心误差。",
    )
    insert_picture_before(conclusion, detection_fig, "图 6 扩展训练集前后的检测性能对比")

    table_caption = insert_paragraph_before(conclusion, "表 15 工作台平面约束后的严格位姿结果")
    table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_caption.paragraph_format.keep_with_next = True
    table_data = [
        ["数据集", "总目标", "位置可用", "法向修正", "严格可用", "严格有效率", "约束后剔除", "跳过"],
        ["训练集", "504", "481", "71", "476", "94.4%", "5", "23"],
        ["验证集", "79", "75", "7", "73", "92.4%", "2", "4"],
        ["测试集", "60", "56", "6", "54", "90.0%", "2", "4"],
    ]
    insert_table_before(conclusion, table_data, [0.7, 0.65, 0.75, 0.75, 0.75, 0.85, 0.85, 0.65])
    insert_picture_before(conclusion, constraint_fig, "图 7 工作台平面约束前后的有效位姿统计")
    insert_picture_before(conclusion, review_fig, "图 8 工作台平面约束后的测试集位姿投影示例", 6.0)
    insert_paragraph_before(
        conclusion,
        "由表 15 可知，测试集 56 个位置可用结果中有 6 个法向受到约束修正，最终 54 个结果同时满足方向偏差不超过 20°、重投影 IoU 不低于 0.45、中心误差不超过 30 px，严格有效率为 90.0%。约束后仍不合格的 2 个目标主要表现为中心偏差或重投影重合度不足，表明平面约束能够修正姿态退化，但不能替代对错误点云分割和严重遮挡的质量筛选。",
    )

    doc.save(str(OUTPUT_DOC))
    print(OUTPUT_DOC)
    print(detection_fig)
    print(constraint_fig)


if __name__ == "__main__":
    main()
